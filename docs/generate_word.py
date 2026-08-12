# -*- coding: utf-8 -*-
"""
Word document generator (APA 7th edition with UNEG adaptations)
for the "Proyecto Saecula" thesis proposal.

Adaptations:
    - 2.54 cm margins on all sides.
    - Justified text.
    - 1.5 line spacing.
APA 7 base format: Times New Roman 12, first-line indent,
hanging indent for references, and page number in the header.

Usage:
    python generate_word.py [input.md] [output.docx]

The document text comes from the Markdown file "proposal.md",
which starts with a metadata block (title page) delimited by "---".
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FUENTE = "Times New Roman"
TAMANO = Pt(12)
INTERLINEADO = 1.5
MARGEN = Cm(2.54)
SANGRIA_PRIMERA = Cm(1.27)   # 0,5 pulgadas (APA 7)
SANGRIA_FRANCESA = Cm(1.27)  # sangría francesa para referencias

ARCHIVO_MD = "proposal.md"
ARCHIVO_SALIDA = "Saecula-Thesis-Proposal.docx"


def aplicar_fuente(run, negrita=False):
    """Aplica Times New Roman de 12 pt a un run, incluyendo alfabetos no latinos."""
    run.font.name = FUENTE
    run.font.size = TAMANO
    run.font.bold = negrita
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for atributo in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(atributo), FUENTE)


def agregar_parrafo(documento, texto="", alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    negrita=False, primera=None, izquierda=None,
                    interlineado=INTERLINEADO, salto_pagina=False,
                    mantener_con_siguiente=False):
    """Agrega un párrafo con el formato base del documento."""
    p = documento.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = alineacion
    pf.line_spacing = interlineado
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if salto_pagina:
        pf.page_break_before = True
    if mantener_con_siguiente:
        pf.keep_with_next = True
    if primera is not None:
        pf.first_line_indent = primera
    if izquierda is not None:
        pf.left_indent = izquierda
    if texto:
        aplicar_fuente(p.add_run(texto), negrita=negrita)
    return p


def agregar_numero_pagina(seccion):
    """Agrega el número de página en la parte superior derecha (APA 7)."""
    encabezado = seccion.header
    p = encabezado.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    aplicar_fuente(run)
    campo = OxmlElement("w:fldChar")
    campo.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = "PAGE"
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    run._r.append(campo)
    run._r.append(instruccion)
    run._r.append(fin)


def leer_metadatos(lineas):
    """Extrae el bloque YAML de la portada delimitado por '---' al inicio."""
    if not lineas or lineas[0].strip() != "---":
        return {}, lineas
    metadatos = {}
    i = 1
    while i < len(lineas) and lineas[i].strip() != "---":
        texto = lineas[i].strip()
        if ":" in texto:
            clave, valor = texto.split(":", 1)
            metadatos[clave.strip()] = valor.strip().strip('"').strip("'")
        i += 1
    return metadatos, lineas[i + 1:]


def dividir_bloques(lineas):
    """Convierte las líneas del cuerpo del Markdown en bloques tipificados."""
    bloques = []
    i = 0
    n = len(lineas)
    while i < n:
        linea = lineas[i].rstrip()
        if not linea.strip():
            i += 1
            continue
        if linea.startswith("### "):
            bloques.append(("h3", linea[4:].strip()))
            i += 1
        elif linea.startswith("## "):
            bloques.append(("h2", linea[3:].strip()))
            i += 1
        elif linea.startswith("# "):
            bloques.append(("h1", linea[2:].strip()))
            i += 1
        elif linea.startswith("- ") or linea.startswith("* "):
            items = []
            while i < n:
                l = lineas[i].rstrip()
                if l.startswith("- ") or l.startswith("* "):
                    items.append(l[2:].strip())
                    i += 1
                elif not l.strip():
                    i += 1
                    break
                else:
                    break
            bloques.append(("bullet", items))
        elif re.match(r"^\d+\.\s", linea):
            items = []
            while i < n:
                l = lineas[i].rstrip()
                m = re.match(r"^\d+\.\s*(.*)$", l)
                if m:
                    items.append(m.group(1).strip())
                    i += 1
                elif not l.strip():
                    i += 1
                    break
                else:
                    break
            bloques.append(("numero", items))
        else:
            parrafo = []
            while i < n:
                l = lineas[i].rstrip()
                if (not l.strip() or l.startswith("#") or l.startswith("- ")
                        or l.startswith("* ") or re.match(r"^\d+\.\s", l)):
                    break
                parrafo.append(l.strip())
                i += 1
            bloques.append(("parrafo", " ".join(parrafo)))
    return bloques


def construir_portada(doc, meta):
    """Construye la portada según el modelo de la UNEG."""
    lineas_portada = [
        ("universidad", False),
        ("vicerrectorado", False),
        ("coordinacion", False),
        ("carrera", False),
    ]
    for clave, negrita in lineas_portada:
        if clave in meta:
            agregar_parrafo(doc, meta[clave], WD_ALIGN_PARAGRAPH.CENTER, negrita)

    for _ in range(3):
        agregar_parrafo(doc)

    titulo = meta.get("titulo", "")
    for parte in titulo.split(". "):
        agregar_parrafo(doc, parte.strip(), WD_ALIGN_PARAGRAPH.CENTER, True)

    for _ in range(3):
        agregar_parrafo(doc)

    if "autores" in meta:
        agregar_parrafo(doc, f"Autores: {meta['autores']}", WD_ALIGN_PARAGRAPH.CENTER)
    if "tutor" in meta:
        agregar_parrafo(doc, f"Tutor(a): {meta['tutor']}", WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        agregar_parrafo(doc)

    lugar = meta.get("lugar", "")
    fecha = meta.get("fecha", "")
    if lugar and fecha:
        agregar_parrafo(doc, f"{lugar}, {fecha}", WD_ALIGN_PARAGRAPH.CENTER)


def construir_cuerpo(doc, bloques):
    """Renderiza los bloques del cuerpo del documento en Word."""
    en_referencias = False

    for tipo, contenido in bloques:
        if tipo == "h1":
            en_referencias = False
            partes = contenido.split(" - ", 1)
            if len(partes) == 2:
                agregar_parrafo(doc, partes[0].upper(), WD_ALIGN_PARAGRAPH.CENTER,
                                True, salto_pagina=True, mantener_con_siguiente=True)
                agregar_parrafo(doc, partes[1].upper(), WD_ALIGN_PARAGRAPH.CENTER,
                                True, mantener_con_siguiente=True)
            else:
                agregar_parrafo(doc, contenido.upper(), WD_ALIGN_PARAGRAPH.CENTER,
                                True, salto_pagina=True, mantener_con_siguiente=True)
        elif tipo == "h2":
            texto = contenido.strip().upper()
            agregar_parrafo(doc, texto, WD_ALIGN_PARAGRAPH.CENTER, True,
                            salto_pagina=(texto == "REFERENCIAS"),
                            mantener_con_siguiente=True)
            en_referencias = texto == "REFERENCIAS"
        elif tipo == "h3":
            agregar_parrafo(doc, contenido, WD_ALIGN_PARAGRAPH.LEFT, True,
                            mantener_con_siguiente=True)
        elif tipo == "parrafo":
            if en_referencias:
                p = agregar_parrafo(doc, contenido, WD_ALIGN_PARAGRAPH.LEFT,
                                    izquierda=SANGRIA_FRANCESA)
                p.paragraph_format.first_line_indent = Cm(-SANGRIA_FRANCESA.cm)
            else:
                agregar_parrafo(doc, contenido, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                primera=SANGRIA_PRIMERA)
        elif tipo == "bullet":
            for item in contenido:
                p = agregar_parrafo(doc, f"• {item}",
                                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                                    izquierda=SANGRIA_PRIMERA)
                p.paragraph_format.first_line_indent = Cm(-0.63)
        elif tipo == "numero":
            for numero, item in enumerate(contenido, start=1):
                p = agregar_parrafo(doc, f"{numero}. {item}",
                                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                                    izquierda=SANGRIA_PRIMERA)
                p.paragraph_format.first_line_indent = Cm(-0.63)


def generar(archivo_md, archivo_salida):
    """Genera el documento Word a partir del archivo Markdown."""
    texto = Path(archivo_md).read_text(encoding="utf-8")
    lineas = texto.splitlines()
    meta, cuerpo = leer_metadatos(lineas)
    bloques = dividir_bloques(cuerpo)

    doc = Document()

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = FUENTE
    estilo_normal.font.size = TAMANO
    estilo_normal.paragraph_format.line_spacing = INTERLINEADO

    seccion = doc.sections[0]
    seccion.top_margin = MARGEN
    seccion.bottom_margin = MARGEN
    seccion.left_margin = MARGEN
    seccion.right_margin = MARGEN
    seccion.page_width = Cm(21.0)   # hoja A4
    seccion.page_height = Cm(29.7)

    agregar_numero_pagina(seccion)

    construir_portada(doc, meta)
    doc.add_page_break()
    construir_cuerpo(doc, bloques)

    doc.core_properties.title = meta.get("titulo", "Propuesta de Trabajo de Grado")
    doc.core_properties.author = meta.get("autores", "")

    doc.save(archivo_salida)
    return archivo_salida


if __name__ == "__main__":
    entrada = sys.argv[1] if len(sys.argv) > 1 else ARCHIVO_MD
    salida = sys.argv[2] if len(sys.argv) > 2 else ARCHIVO_SALIDA
    ruta = generar(entrada, salida)
    print(f"Documento generado correctamente: {ruta}")
